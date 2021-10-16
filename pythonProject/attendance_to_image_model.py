# -*- coding: utf-8 -*-

import os.path
from pdf2image import convert_from_bytes
import copy
import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_path(path):
    if not os.path.exists(path):
        os.mkdir(path)


def convert_to_pdf(sheet, teacher_list, save_path):
    value = sheet.get_all_values(include_tailing_empty_rows=False)
    teacher_info = []  # contains cell number of teacher ([teacher, starting_row, ending_row, attendance[]])
    for row in range(len(value)):
        if value[row][0] in teacher_list:
            teacher = value[row][0]
            attendance = [] # attendance[(student, count, pay)]
            starting_row = row + 1
            for i in range(row, len(value)):
                if value[i][0] == '' and value[i][1] != '':
                    d = copy.deepcopy(value[i][4:14])
                    while ('' in d):
                        d.remove('')
                    count = len(d)
                    if value[i][3] == '':
                        pay = '미납'
                    else:
                        pay = '완료'
                    if not ((count < 3 and pay == '완료') or (count == 0 and pay == '미납')):
                        attendance.append((value[i][1], count, pay))

                elif i + 1 < len(value) and (value[i + 1][0] != '' or value[i + 1][0] != ""):
                    ending_row = i
                    teacher_info.append([teacher, starting_row, ending_row, attendance])
                    break
                elif i + 1 == len(value):
                    ending_row = len(value)
                    teacher_info.append([teacher, starting_row, ending_row, attendance])
                    break
    return teacher_info



    # for teacher_data in teacher_info:
    #     range_address = '&c1=1' + '&r1=' + str(teacher_data[1] - 1) + '&c2=14' + '&r2=' + str(teacher_data[2] - 1)
    #     url = ('https://docs.google.com/spreadsheets/d/' + '1VM00nved7joz20rRD_M1k_E6F0ZFCG70RQ5KfjRzcgY' + '/export?'
    #            + 'format=pdf'  # export as PDF
    #            + '&portrait=false'  # landscape
    #            + '&top_margin=0.00'  # Margin
    #            + '&bottom_margin=0.00'  # Margin
    #            + '&left_margin=0.00'  # Margin
    #            + '&right_margin=0.00'  # Margin
    #            + '&pagenum=RIGHT'  # Put page number to right of footer
    #            + '&gid=' + str(sheet.id)
    #            + range_address)
    #
    #     pdf_name = '\\' + teacher[0] + '.pdf'
    #     r = requests.get(url, allow_redirects=True)
    #     with open(save_path + pdf_name, 'wb') as f:
    #         f.write(r.content)


def save_to_doc(teacher_info):
    date = datetime.datetime.today()
    document = Document()


    section = document.sections[0]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '3')

    style = document.styles['Normal']
    font = style.font

    paragraph = document.add_paragraph()
    run = paragraph.add_run(str(date.year) + '-' + str(date.month) + '-' + str(date.day) + ' 출결 현황\n')
    run.bold = True
    paragraph.paragraph_format.space_after = 0

    for teacher_data in teacher_info:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = 0
        run = paragraph.add_run(teacher_data[0] + ' 강사님')
        run.bold = True
        for i in teacher_data[3]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(i[0] + ' ' + str(i[1]) + '회 ' + i[2])
            paragraph.paragraph_format.space_after = 0
        paragraph = document.add_paragraph('')
        paragraph.paragraph_format.space_after = 0


    document.save(r'C:\Users\admin\Desktop\출결 현황.docx')

def convert_pdf_to_image(save_path):
    file_list = os.listdir(save_path)
    for file_name in file_list:
        image_path = save_path + '\\' + str(file_name)[:-4]
        print(image_path)
        create_path(image_path)
        pages = convert_from_bytes(open(save_path + '\\' + file_name, 'rb').read(),
                                   poppler_path=r'C:\Program Files\poppler-21.03.0\Library\bin')
        for i, page in enumerate(pages):
            page.save(image_path + '\\' + str(file_name)[:-4] + str(i) + '.jpg', 'JPEG')

